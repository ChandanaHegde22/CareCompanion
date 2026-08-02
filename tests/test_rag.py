"""tests/test_rag.py – RAG document processing and retrieval tests."""
import os, pytest, tempfile, uuid


def uid():
    from services.auth_service import register_user
    u = f"rag_{uuid.uuid4().hex[:8]}"
    r = register_user(u, f"{u}@ex.com", "Pass1234!")
    return r["user_id"]


class TestDocumentProcessor:
    def test_extract_txt(self, tmp_path):
        from rag.document_processor import extract_text_txt
        f = tmp_path / "sample.txt"
        f.write_text("Blood pressure: 120/80. HbA1c: 7.2%.")
        text = extract_text_txt(str(f))
        assert "Blood pressure" in text and "HbA1c" in text

    def test_chunk_produces_multiple_chunks(self):
        from rag.document_processor import chunk_text
        text = "This is a sentence. " * 100
        chunks = chunk_text(text, "test.txt", chunk_size=200, overlap=20)
        assert len(chunks) > 1

    def test_chunks_have_required_keys(self):
        from rag.document_processor import chunk_text
        chunks = chunk_text("Word " * 200, "report.txt", chunk_size=150)
        assert all("text" in c and "source" in c and "chunk_index" in c for c in chunks)

    def test_chunk_index_sequential(self):
        from rag.document_processor import chunk_text
        chunks = chunk_text("Word " * 300, "test.txt", chunk_size=100, overlap=10)
        assert [c["chunk_index"] for c in chunks] == list(range(len(chunks)))

    def test_chunk_source_matches(self):
        from rag.document_processor import chunk_text
        chunks = chunk_text("Medical content. " * 40, "blood_test.pdf")
        assert all(c["source"] == "blood_test.pdf" for c in chunks)

    def test_clean_text_no_double_spaces(self):
        from rag.document_processor import clean_text
        clean = clean_text("This  has   extra   spaces")
        assert "  " not in clean

    def test_clean_text_no_triple_newlines(self):
        from rag.document_processor import clean_text
        clean = clean_text("Line\n\n\n\nBreak")
        assert "\n\n\n" not in clean

    def test_process_txt_file(self, tmp_path):
        from rag.document_processor import process_document
        f = tmp_path / "report.txt"
        f.write_text("Patient: John Doe. Diagnosis: Hypertension. BP: 150/90. " * 15)
        chunks = process_document(str(f), "report.txt")
        assert len(chunks) > 0

    def test_empty_file_returns_no_chunks(self, tmp_path):
        from rag.document_processor import process_document
        f = tmp_path / "empty.txt"
        f.write_text("")
        assert process_document(str(f), "empty.txt") == []

    def test_process_docx(self, tmp_path):
        try:
            from docx import Document
            from rag.document_processor import extract_text_docx
            doc = Document()
            doc.add_paragraph("Patient: Ramesh Kumar")
            doc.add_paragraph("Diagnosis: Type 2 Diabetes")
            p = str(tmp_path / "report.docx")
            doc.save(p)
            text = extract_text_docx(p)
            assert "Ramesh Kumar" in text
        except ImportError:
            pytest.skip("python-docx not installed")


class TestEmbedder:
    def test_single_embedding_shape(self):
        try:
            from rag.embedder import get_embedding, embedding_dimension
            emb = get_embedding("blood pressure 120/80")
            if emb is None:
                pytest.skip("Embedder model not loaded")
            assert emb.shape == (embedding_dimension(),)
        except ImportError:
            pytest.skip("sentence-transformers not installed")

    def test_batch_embeddings_shape(self):
        try:
            from rag.embedder import get_embeddings, embedding_dimension
            texts = ["medicine A", "medicine B", "blood test"]
            embs  = get_embeddings(texts)
            if embs is None:
                pytest.skip("Embedder not available")
            assert embs.shape == (3, embedding_dimension())
        except ImportError:
            pytest.skip("sentence-transformers not installed")

    def test_embeddings_normalized(self):
        try:
            import numpy as np
            from rag.embedder import get_embedding
            emb = get_embedding("normalize this text")
            if emb is None:
                pytest.skip("Embedder not available")
            norm = float(np.linalg.norm(emb))
            assert abs(norm - 1.0) < 0.05
        except ImportError:
            pytest.skip("sentence-transformers not installed")


class TestRetriever:
    CHUNKS = [
        {"text": "Patient blood pressure is 140/90 mmHg.", "source": "bp.pdf", "chunk_index": 0},
        {"text": "HbA1c level is 7.2%. Diabetes well controlled.", "source": "blood.pdf", "chunk_index": 0},
        {"text": "Prescribed Metformin 500mg twice daily after food.", "source": "rx.pdf", "chunk_index": 0},
        {"text": "Cholesterol total 180 mg/dl. HDL 55, LDL 100.", "source": "lipid.pdf", "chunk_index": 0},
        {"text": "Follow-up visit on 15th January 2025.", "source": "discharge.pdf", "chunk_index": 0},
    ]

    def test_no_index_initially(self):
        from rag.retriever import has_index
        assert has_index(99999) is False

    def test_build_and_search(self):
        try:
            from rag.retriever import build_and_save_index, search, has_index
            user_id = uid()
            ok = build_and_save_index(user_id, self.CHUNKS)
            if not ok:
                pytest.skip("FAISS/embedder not available")
            assert has_index(user_id)
            results = search(user_id, "blood pressure", top_k=2)
            assert len(results) > 0
            texts = " ".join(r["text"] for r in results)
            assert "140/90" in texts or "blood pressure" in texts.lower()
        except ImportError:
            pytest.skip("faiss-cpu or sentence-transformers not installed")

    def test_results_have_scores(self):
        try:
            from rag.retriever import build_and_save_index, search
            user_id = uid()
            ok = build_and_save_index(user_id, self.CHUNKS)
            if not ok:
                pytest.skip("FAISS not available")
            results = search(user_id, "diabetes hba1c", top_k=3)
            assert all(isinstance(r["score"], float) for r in results)
        except ImportError:
            pytest.skip("faiss-cpu not installed")

    def test_index_size(self):
        try:
            from rag.retriever import build_and_save_index, get_index_size
            user_id = uid()
            ok = build_and_save_index(user_id, self.CHUNKS)
            if not ok:
                pytest.skip("FAISS not available")
            assert get_index_size(user_id) == len(self.CHUNKS)
        except ImportError:
            pytest.skip("faiss-cpu not installed")


class TestRAGChain:
    def test_no_documents_returns_guidance(self):
        from rag.rag_chain import answer_medical_query
        user_id = uid()
        result  = answer_medical_query(user_id, "What is my blood pressure?")
        assert result["success"] is False
        answer_lower = result["answer"].lower()
        assert "upload" in answer_lower or "document" in answer_lower or "report" in answer_lower

    def test_document_list_is_list(self):
        from rag.rag_chain import get_user_documents
        user_id = uid()
        assert isinstance(get_user_documents(user_id), list)

    def test_answer_empty_query(self):
        from rag.rag_chain import answer_medical_query
        user_id = uid()
        result  = answer_medical_query(user_id, "   ")
        assert result["success"] is False
